
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sqlalchemy import text
from collections import defaultdict

import os

from config import REPORT_FILE_NAME


# ----------------------------------------------------
# INSERT IMAGES 2 PER ROW
# ----------------------------------------------------

def insert_images_two_per_row(cell, image_paths):

    if not image_paths:
        return

    cell.text = ""

    rows = (len(image_paths) + 1) // 2
    img_table = cell.add_table(rows=rows, cols=2)

    index = 0

    for r in range(rows):
        for c in range(2):

            if index >= len(image_paths):
                break

            img_path = image_paths[index]

            try:

                paragraph = img_table.cell(r, c).paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = paragraph.add_run()

                run.add_picture(
                    img_path,
                    width=Cm(3.48),
                    height=Cm(1.96)
                )

            except Exception:
                pass

            index += 1


# ----------------------------------------------------
# TABLE BORDER FUNCTION
# ----------------------------------------------------

def set_cell_border(cell):

    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('top','left','bottom','right'):

        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '8')
        element.set(qn('w:color'), '000000')

        tcBorders.append(element)

    tcPr.append(tcBorders)



# ----------------------------------------------------
# MAIN REPORT GENERATOR
# ----------------------------------------------------

def generate_docx_report(session, audit_id, image_folder):

    document = Document()

    # ------------------------------------------------
    # POTRAIT PAGE
    # ------------------------------------------------

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT

    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    document.add_heading("Electrical Safety Audit Report", level=0)
    document.add_paragraph("")


    faults = session.execute(
        text("""
        SELECT cluster_id, building, fault_type, message, image_path
        FROM faults
        WHERE audit_id = :audit_id
        ORDER BY cluster_id
        """),
        {"audit_id": audit_id}
    ).fetchall()

    if not faults:
        document.add_paragraph("No faults recorded.")
        document.save(REPORT_FILE_NAME)
        print("Empty report generated.")
        return


    # ------------------------------------------------
    # GROUP FAULTS
    # ------------------------------------------------

    grouped_faults = defaultdict(lambda: {
        "fault_type": set(),
        "remarks": [],
        "locations": set(),
        "images": set()
    })

    for fault in faults:

        cluster = fault.cluster_id

        grouped_faults[cluster]["fault_type"].add(fault.fault_type)
        grouped_faults[cluster]["remarks"].append(fault.message)
        grouped_faults[cluster]["locations"].add(fault.building)

        if fault.image_path:

            full_path = os.path.join(image_folder, fault.image_path)

            if os.path.exists(full_path):
                grouped_faults[cluster]["images"].add(full_path)


    # ------------------------------------------------
    # CREATE TABLE
    # ------------------------------------------------

    
    table = document.add_table(rows=1, cols=3)
    table.autofit = False

    widths = [1.38, 7.82, 7.92]

    tbl = table._element
    tblPr = tbl.tblPr

    # force fixed layout
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)

    # define column grid
    tblGrid = tbl.tblGrid

    # clear existing grid
    for child in list(tblGrid):
        tblGrid.remove(child)

    for w in widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w * 567)))   # cm → Word units
        tblGrid.append(gridCol)

    # also apply width to cells
    for i,w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    header = table.rows[0]

    header.height = Cm(1.04)
    header.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    header.cells[0].text = "Si No"
    header.cells[1].text = "Images"
    header.cells[2].text = "Remarks"

    for cell in header.cells:
        set_cell_border(cell)


    
    # ---------------------------------------------------------
    # ADD DATA
    # ---------------------------------------------------------
    serial_number = 1

    for cluster, data in grouped_faults.items():

        print("Cluster", cluster)
        print("Cluster images:", data["images"])

        row_cells = table.add_row().cells

        row_cells[0].text = str(serial_number)

        images = list(data["images"])[:20]

        insert_images_two_per_row(row_cells[1], images)

        locations = ", ".join(sorted(data["locations"]))
        remarks_list = list(dict.fromkeys(data["remarks"]))[:20]
        fault_types = ", ".join(sorted(data["fault_type"]))

        cell = row_cells[2]

        # first line
        p = cell.paragraphs[0]
        p.text = f"\n{fault_types}"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # remarks
        for r in remarks_list:
            bullet = cell.add_paragraph()
            bullet.text = f"- {r}"
            bullet.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # locations
        loc = cell.add_paragraph()
        loc.text = f"\n{locations}"
        loc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        

        serial_number += 1


    document.save(REPORT_FILE_NAME)

    print(f"Report generated successfully: {REPORT_FILE_NAME}")
